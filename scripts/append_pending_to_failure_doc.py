"""Append the 'Pending Tasks' and 'Resolution Log' sections to the failure analysis docx.

Standalone script; not part of the package. Run from repo root.
"""
from docx import Document


def main() -> None:
    doc = Document("Trading_System_Failure_Analysis_and_Summary.docx")

    # ------------------------------------------------------------------ #
    # 6. Pending Tasks (Updated 2026-08-29)
    # ------------------------------------------------------------------ #
    doc.add_heading("6. Pending Tasks (Updated 2026-08-29)", level=1)
    doc.add_paragraph(
        "A spec/plan/tasks/checklist set has been generated for the news-driven GATv2 GNN "
        "retrain under specs/003-news-driven-gnn-retrain/. The items below are the live "
        "punch list derived from the spec, in execution order. Items are removed (not "
        "checked) when shipped; the live task list is the source of truth."
    )

    doc.add_heading("6.1 Pre-Monday (Sat 29 Aug to Sun 30 Aug 2026)", level=2)
    pre = [
        ("T001-T006",
         "Setup: add sql/52_news_snapshot.sql, sql/53_gnn_news_columns.sql, "
         "MessageType.RESEARCH_VIEW, ResearchOutput Pydantic model, "
         "config/agents.yaml research.enabled default false, config/gnn.yaml news edge reasons."),
        ("T007-T012",
         "Foundational: built-in finance lexicon (~200 words), Alpaca fetch_news wrapper "
         "with retry-with-backoff, backfill CLI scaffold, tests for schema/lexicon/protocol."),
        ("T013-T022",
         "US1: research node + backfill CLI + news_snapshot writes + feature-flag-off "
         "no-op + decision-time news_cutoff param. End state: news table populated, research node live."),
        ("T023-T035",
         "US2: extend option_graph.py with news_cooccurrence and news_sentiment_correlation "
         "edge types, 8-d edge feature layout, lifting rule from underlying to contract pairs, "
         "decision-time mask. Backfill 30+ snapshots."),
        ("T036-T049",
         "US3: walk_forward_ablation trainer with 3 folds x 4 configs (xgb / gcn_v1 / "
         "gatv2_prenews / gatv2_news), sharpe_approx, promote_default rule, reproducibility "
         "with fixed seed. Run it. Decide on promotion."),
        ("T050-T061",
         "US4: wire research node into orchestrator graph (volatility -> research -> "
         "options_structure), InferenceService.build_snapshot news_block, supervisor "
         "news_sentiment field, env-var override, graceful fallback on missing artifact / "
         "unreachable Alpaca / sparse news."),
    ]
    for tid, desc in pre:
        p = doc.add_paragraph()
        p.add_run(f"{tid}: ").bold = True
        p.add_run(desc)

    doc.add_heading("6.2 Monday Open (Mon 31 Aug 13:30 UTC)", level=2)
    mon = [
        ("T062-T070",
         "Polish: update _post_fix_rating.md, phase-2-gnn/improvements.md, README/NEWS docs; "
         "eod_pnl notes carries news-flag state; run full quickstart Phases A-H; pytest -q "
         "green; 5-min load test confirms SC-005/SC-006."),
        ("Promotion gate",
         "Flip agents.research.enabled to true ONLY if models/ablation_option_opportunity.json "
         "promote_default.model == 'gatv2_news' AND promote_default.promote == true. Otherwise "
         "leave the flag off and demo with the pre-news baseline."),
        ("Kill switch",
         "Per the original Section 2.5 of this document: a manual pause/halt-all-trading control "
         "must be built and tested before live trading starts. This is unaffected by the news "
         "feature but is a hard pre-condition for Monday."),
    ]
    for tid, desc in mon:
        p = doc.add_paragraph()
        p.add_run(f"{tid}: ").bold = True
        p.add_run(desc)

    doc.add_heading("6.3 Risks accepted by this plan", level=2)
    risks = [
        "News edges may NOT beat the pre-news GATv2 baseline on the held-out fold. The trainer "
        "is feature-flagged default-off; we only flip the flag on for live trading if the "
        "ablation recommends it. The pre-news baseline ships in any case.",
        "Alpaca news API may not have full coverage for all 15 universe symbols. The pipeline "
        "still runs; news edges are simply absent for pairs with no co-mentions.",
        "Loughran-McDonald CSV not shipped in the repo. A built-in ~200-word finance list ships "
        "as the fallback; FR-014 forbids LLM-based sentiment.",
        "Pre-existing GATv2 edge_dim=8 vs option_graph 10-d edge-features mismatch. Resolved in "
        "US2 (T030) by changing the one-hot from 4-d to 2-d; total stays at 8. This is the only "
        "schema-affecting change.",
        "GPU is unavailable in CI. The GATv2 model trains on a single CPU in < 5 minutes; we "
        "do not need a GPU.",
    ]
    for r in risks:
        doc.add_paragraph(f"  - {r}")

    doc.add_heading("6.4 Cross-references", level=2)
    xrefs = [
        "specs/003-news-driven-gnn-retrain/spec.md - 4 user stories, 14 functional requirements, 6 success criteria",
        "specs/003-news-driven-gnn-retrain/plan.md - technical context, constitution check, project structure",
        "specs/003-news-driven-gnn-retrain/research.md - 8 decisions resolving every NEEDS CLARIFICATION",
        "specs/003-news-driven-gnn-retrain/data-model.md - entity definitions and SQL DDL",
        "specs/003-news-driven-gnn-retrain/contracts/ - JSON-Schema for news_snapshot, ResearchOutput, AblationRow",
        "specs/003-news-driven-gnn-retrain/quickstart.md - runnable validation steps for every acceptance scenario",
        "specs/003-news-driven-gnn-retrain/tasks.md - 70 tasks across 7 phases",
        "specs/003-news-driven-gnn-retrain/checklists/requirements.md - spec quality checklist",
        "specs/003-news-driven-gnn-retrain/checklists/requirements-quality.md - reviewer-owned requirements-quality checklist (54 items)",
    ]
    for x in xrefs:
        doc.add_paragraph(f"  - {x}")

    doc.add_paragraph(
        "This section is regenerated by the speckit workflow on every spec/plan/tasks re-run; "
        "in between, task status lives in the IDE task list and the spec/003 tasks.md file. "
        "The single source of truth is tasks.md."
    )

    # ------------------------------------------------------------------ #
    # 7. Resolution Log (running, appended)
    # ------------------------------------------------------------------ #
    doc.add_heading("7. Resolution Log (running, appended)", level=1)
    doc.add_paragraph(
        "Each entry below records a failure-point close-out action. The format is "
        "'[date] Section N.M - <action taken> - <artifact path>'. Newest entries last."
    )
    resolutions = [
        ("2026-08-29", "Section 2.2 (row 1: Overfitting on a short history)",
         "Adopted 3 expanding-window folds; held-out fold 2 is the validation slice. "
         "Implemented in tasks.md T042.", "specs/003-news-driven-gnn-retrain/tasks.md"),
        ("2026-08-29", "Section 2.2 (row 4: News embedding pipeline unfinished by Monday)",
         "Adopted feature-flag default-off + walk-forward promotion rule. Implemented in "
         "spec.md FR-002, FR-009, FR-010.", "specs/003-news-driven-gnn-retrain/spec.md"),
        ("2026-08-29", "Section 2.1 (row 2: News timing leakage)",
         "Adopted news_cutoff parameter on the option-graph builder; backtest path masks "
         "articles published after the decision timestamp. Implemented in spec.md FR-003 "
         "and tasks.md T033.", "specs/003-news-driven-gnn-retrain/spec.md"),
        ("2026-08-29", "Section 3 (row 4: News feature flag)",
         "Spec covers the flag (FR-002), the promotion rule (FR-009), and the graceful "
         "fallback (FR-010). Tasks T050-T054 cover the test surface; T055-T061 cover the "
         "implementation.", "specs/003-news-driven-gnn-retrain/tasks.md"),
        ("2026-08-29", "Section 3 (row 5: A2A decision logs)",
         "No change required by spec 003; the existing journal structure already captures "
         "agent_messages. Spec 003 adds the research node to the same journal row, not a "
         "new table.", "src/agents/journal.py (existing)"),
        ("2026-08-29", "Section 3 (row 11: Demo script)",
         "Quickstart Phase H documents the demo_script CLI; tasks T066 covers the runnable "
         "validation.", "specs/003-news-driven-gnn-retrain/quickstart.md"),
    ]
    for d, sec, action, art in resolutions:
        p = doc.add_paragraph()
        p.add_run(f"[{d}] {sec}: ").bold = True
        p.add_run(f"{action} Artifact: {art}")

    doc.add_paragraph(
        "This Resolution Log is appended to (never trimmed) on every spec/plan/tasks re-run. "
        "When all of Sections 2 and 3 are closed out and the Monday-Friday hackathon window "
        "is complete, this document can be retired in favor of a final post-mortem."
    )

    doc.save("Trading_System_Failure_Analysis_and_Summary.docx")
    print(f"OK, paragraphs now: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
